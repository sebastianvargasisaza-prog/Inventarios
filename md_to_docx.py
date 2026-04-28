"""
Convierte los documentos markdown de Alejandro a .docx con formato profesional.

Uso: python md_to_docx.py
Genera:
  - BROCHURE_CORTEX_LABS.docx
  - DOCUMENTO_ALEJANDRO_SAAS.docx
  - DOCUMENTO_ALEJANDRO_FACTURACION.docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===================== Helpers =====================

def _strip_emoji(s: str) -> str:
    """Quita emojis comunes que rompen render en Word si la fuente no los soporta."""
    # Mantenemos emojis simples; solo limpiamos los que dan problema o decorativos al inicio.
    return s

def _shade_cell(cell, color_hex):
    """Pone color de fondo a una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def _set_cell_borders(cell, color="999999", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tc_pr.append(tcBorders)


# ===================== Inline parsing =====================

def add_inline_runs(paragraph, text, base_bold=False, base_color=None, base_size=None):
    """Procesa texto inline con **bold**, *italic*, `code`, [link](url)."""
    # Patrón: bold, italic, code, link
    pattern = re.compile(
        r'\*\*(?P<bold>.+?)\*\*'
        r'|\*(?P<italic>.+?)\*'
        r'|`(?P<code>.+?)`'
        r'|\[(?P<link_text>.+?)\]\((?P<link_url>.+?)\)'
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.bold = base_bold
            if base_color:
                run.font.color.rgb = base_color
            if base_size:
                run.font.size = base_size
        if m.group('bold') is not None:
            run = paragraph.add_run(m.group('bold'))
            run.bold = True
            if base_color: run.font.color.rgb = base_color
            if base_size: run.font.size = base_size
        elif m.group('italic') is not None:
            run = paragraph.add_run(m.group('italic'))
            run.italic = True
            if base_color: run.font.color.rgb = base_color
            if base_size: run.font.size = base_size
        elif m.group('code') is not None:
            run = paragraph.add_run(m.group('code'))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif m.group('link_text') is not None:
            run = paragraph.add_run(m.group('link_text'))
            run.font.color.rgb = RGBColor(0x10, 0x66, 0xCC)
            run.underline = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold
        if base_color: run.font.color.rgb = base_color
        if base_size: run.font.size = base_size


# ===================== Markdown → docx =====================

def md_to_docx(md_path: Path, docx_path: Path, brand_color=RGBColor(0x6D, 0x28, 0xD9)):
    """Convierte un .md a .docx con formato profesional."""
    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    doc = Document()

    # Estilo base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Línea vacía
        if not stripped:
            i += 1
            continue

        # Separador horizontal ---
        if stripped == '---':
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            p_pr.append(pBdr)
            i += 1
            continue

        # Títulos
        if stripped.startswith('# '):
            h = doc.add_paragraph()
            run = h.add_run(_strip_emoji(stripped[2:]))
            run.bold = True
            run.font.size = Pt(28)
            run.font.color.rgb = brand_color
            h.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        if stripped.startswith('## '):
            h = doc.add_paragraph()
            run = h.add_run(_strip_emoji(stripped[3:]))
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = brand_color
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        if stripped.startswith('### '):
            h = doc.add_paragraph()
            run = h.add_run(_strip_emoji(stripped[4:]))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        if stripped.startswith('#### '):
            h = doc.add_paragraph()
            run = h.add_run(_strip_emoji(stripped[5:]))
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Bloque de código ```
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # saltar ```
            p = doc.add_paragraph()
            _shade_para(p, 'F3F4F6')
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        # Cita >
        if stripped.startswith('> '):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                line_q = lines[i].strip()
                quote_lines.append(line_q[1:].strip() if line_q.startswith('>') else line_q)
                i += 1
            quote_text = ' '.join(l for l in quote_lines if l)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            _shade_para(p, 'F5F3FF')
            add_inline_runs(p, quote_text, base_color=RGBColor(0x4C, 0x1D, 0x95))
            continue

        # Tabla
        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'\|[\s\-:|]+\|', lines[i+1].strip()):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            _add_md_table(doc, table_lines, brand_color)
            continue

        # Lista bullet
        if stripped.startswith('- ') or stripped.startswith('* '):
            list_items = []
            while i < len(lines):
                ls = lines[i].strip()
                if ls.startswith('- ') or ls.startswith('* '):
                    list_items.append(ls[2:])
                    i += 1
                elif ls == '':
                    i += 1
                    if i < len(lines) and not (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                        break
                else:
                    break
            for item in list_items:
                p = doc.add_paragraph(style='List Bullet')
                add_inline_runs(p, item)
            continue

        # Lista numerada
        if re.match(r'^\d+\.\s', stripped):
            list_items = []
            while i < len(lines):
                ls = lines[i].strip()
                m = re.match(r'^\d+\.\s(.+)', ls)
                if m:
                    list_items.append(m.group(1))
                    i += 1
                elif ls == '':
                    i += 1
                    if i < len(lines) and not re.match(r'^\d+\.\s', lines[i].strip()):
                        break
                else:
                    break
            for item in list_items:
                p = doc.add_paragraph(style='List Number')
                add_inline_runs(p, item)
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f'  -> {docx_path.name} OK')


def _shade_para(paragraph, color_hex):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    p_pr.append(shd)


def _add_md_table(doc, table_lines, brand_color):
    """Convierte líneas markdown |a|b|c| a tabla docx."""
    # Quita la línea de separador
    rows = []
    for ln in table_lines:
        if re.match(r'^\|[\s\-:|]+\|$', ln):
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < n_cols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Header
    for ci, cell_text in enumerate(rows[0]):
        cell = table.cell(0, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        add_inline_runs(p, cell_text, base_bold=True, base_color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(cell, '6D28D9')

    # Body
    for ri in range(1, len(rows)):
        for ci, cell_text in enumerate(rows[ri]):
            cell = table.cell(ri, ci)
            cell.text = ''
            add_inline_runs(cell.paragraphs[0], cell_text)
            if ri % 2 == 0:
                _shade_cell(cell, 'F9FAFB')


# ===================== Main =====================

if __name__ == '__main__':
    base = Path(__file__).parent
    targets = [
        ('BROCHURE_CORTEX_LABS.md',           'BROCHURE_CORTEX_LABS.docx',           RGBColor(0x6D, 0x28, 0xD9)),  # violeta
        ('DOCUMENTO_ALEJANDRO_SAAS.md',       'DOCUMENTO_ALEJANDRO_SAAS.docx',       RGBColor(0x0E, 0x76, 0x90)),  # teal
        ('DOCUMENTO_ALEJANDRO_FACTURACION.md','DOCUMENTO_ALEJANDRO_FACTURACION.docx',RGBColor(0xB4, 0x53, 0x09)),  # ambar oscuro
        ('GAP_ENTERPRISE_VS_ODOO_SAP.md',     'GAP_ENTERPRISE_VS_ODOO_SAP.docx',     RGBColor(0xC0, 0x37, 0x44)),  # rojo carmesi
        ('ROADMAP_MOVIL_NATIVO.md',           'ROADMAP_MOVIL_NATIVO.docx',           RGBColor(0x16, 0xA3, 0x4A)),  # verde
    ]
    for src, dst, color in targets:
        src_path = base / src
        dst_path = base / dst
        if not src_path.exists():
            print(f'  ! {src} no existe')
            continue
        try:
            print(f'Convirtiendo {src} ...')
            md_to_docx(src_path, dst_path, brand_color=color)
        except PermissionError:
            print(f'  ! {dst} esta abierto en Word, saltando (cierra Word para regenerar)')
        except Exception as e:
            print(f'  ! Error en {src}: {e}')
    print('\nListo. Archivos .docx en:', base)
